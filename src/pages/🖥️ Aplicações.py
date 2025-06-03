from helper_methods import (
    initialize_application,
    show_sidebar,
    encrypt_password,
    decrypt_password,
    PROJECTS_DIR,
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as langchain_Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from helper_spingestion import SharePointDownloader
from typing import Tuple, Dict
from datetime import datetime
from docx import Document
import streamlit as st
import pandas as pd
import pdfplumber
import tempfile
import zipfile
import shutil
import stat
import time
import json
import re
import os

# Setup
initialize_application()
show_sidebar()

# Constants
FLAG_SAVE_FILE = False

# Ensure the projects directory exists
os.makedirs(PROJECTS_DIR, exist_ok=True)


def on_rm_error(func, path, exc_info):
    """Error handler for read-only files."""
    os.chmod(path, stat.S_IWRITE)
    func(path)


def validate_project_inputs(
    project_name, project_owner, password, prompt_app, projects_dir
):
    """
    Validates the inputs for creating or updating a project.
    Args:
        project_name (str): Name of the project.
        project_owner (str): Owner of the project.
        password (str): Password for the project.
        prompt_app (str): Prompt for the application.
        projects_dir (str): Directory where projects are stored.

        Returns:
            Tuple[bool, str]: A tuple containing a boolean indicating success or failure,
            and a string with an error message if validation fails.
    """
    if not project_name:
        return False, "Por favor, digite o nome do projeto."
    if len(project_name) >= 30:
        return False, "O nome do projeto deve ter no máximo 30 caracteres."
    if not project_owner:
        return False, "Por favor, digite seu usuário."
    if len(project_owner) >= 30:
        return False, "O nome do usuário deve ter no máximo 30 caracteres."
    if os.path.exists(os.path.join(projects_dir, project_name)):
        return False, "Já existe um projeto com esse nome."
    if password and len(password) < 5:
        return False, "A senha deve ter no mínimo 5 caracteres."
    if password and len(password) >= 25:
        return False, "A senha deve ter no máximo 25 caracteres."
    if prompt_app and len(prompt_app) < 50:
        return False, "O prompt deve ter mais de 50 caracteres."
    if prompt_app and len(prompt_app) >= 2_000:
        return False, "O prompt deve ter no máximo 2000 caracteres."
    return True, None


# Function to extract text and word count from various file formats
def extract_text_and_word_count(file_path: str) -> Tuple[str, int]:
    """
    Extracts text content and word count from various file formats.

    Args:
        file_path (str): Path to the input file.

    Returns:
        Tuple[str, int]: Extracted text content and the corresponding word count.
    """
    text = ""
    try:
        if file_path.endswith(".pdf"):
            with pdfplumber.open(file_path) as pdf:
                text = "\n".join(
                    [page.extract_text() for page in pdf.pages if page.extract_text()]
                )
        elif file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        elif file_path.endswith(".csv"):
            df = pd.read_csv(file_path, encoding="utf-8")
            text = df.to_string()
        elif file_path.endswith(".xlsx"):
            df = pd.read_excel(file_path, engine="openpyxl")
            text = df.to_string()
        elif file_path.endswith(".docx"):
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif file_path.endswith(".json"):
            with open(file_path, "r", encoding="utf-8") as f:
                text = json.dumps(json.load(f))

        # check text
        if text and len(text) > 0:
            text = re.sub(r"\s+", " ", text).strip()
            return text, len(text)
        else:
            return None, 0

    except Exception:
        return None, 0


def save_context_database(
    project_name: str, dict_projects_text: Dict[str, str], total_word_count: int
) -> None:
    """
    Saves extracted text data either as a plain text file (if word count is small) or as a FAISS vector database.

    Args:
        project_name (str): Name of the project.
        dict_projects_text (Dict[str, str]): Dictionary containing filenames and their corresponding extracted text.
        total_word_count (int): Total word count of all extracted text.
    """
    parameters = st.session_state["parameters"]
    project_dir = os.path.join(PROJECTS_DIR, project_name)

    # KB
    if (
        len(dict_projects_text) == 1
        and total_word_count <= parameters["max_word_count_kb"]
    ):
        kb_path = os.path.join(project_dir, "KB.txt")
        with open(kb_path, "w", encoding="utf-8") as file:
            all_text = list(dict_projects_text.values())[0]
            file.write("\n".join(all_text))

    # Context (vector DB)
    else:
        faiss_path = os.path.join(project_dir, "faiss_db")
        os.makedirs(faiss_path, exist_ok=True)

        # text splitter
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=parameters["chunk_size"],
            chunk_overlap=parameters["chunk_overlap"],
            separators=["\n\n", "\n", ".", "!", "?", ","],
        )

        # Iterate dictionary
        split_documents = []
        for filename, text in dict_projects_text.items():
            chunks = text_splitter.split_text(text)
            for chunk in chunks:
                doc = langchain_Document(
                    page_content=chunk, metadata={"source": os.path.basename(filename)}
                )
                split_documents.append(doc)

        # embedding model
        embedding_model = HuggingFaceEmbeddings(
            model_name=parameters["embedding"],
            encode_kwargs={"normalize_embeddings": False},
        )

        vectorstore = FAISS.from_documents(split_documents, embedding_model)
        vectorstore.save_local(faiss_path)


def process_files_and_store_embeddings(
    project_name: str, project_owner: str, password: str
) -> None:
    """
    Processes uploaded files, extracts content, and stores embeddings.

    Args:
        project_name (str): Name of the project.
        project_owner (str): Owner of the project.
    """
    project_path = os.path.join(PROJECTS_DIR, project_name)
    files_path = os.path.join(project_path, "files")
    metadata_path = os.path.join(project_path, "metadata.json")

    # create dir
    os.makedirs(project_path, exist_ok=True)

    # flag password
    if password == "":
        flag_password = False
    else:
        flag_password = True

    try:
        metadata: Dict[str, any] = {
            "project_name": project_name,
            "project_owner": project_owner,
            "flag_password": flag_password,
            "password": encrypt_password(password),
            "creation_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "files": [],
        }

        total_word_count = 0
        qtd_projects = 0
        dict_projects_text = {}
        for file_name in os.listdir(files_path):
            file_path = os.path.join(files_path, file_name)
            text, word_count = extract_text_and_word_count(file_path)
            if text is None or len(text) <= 5:
                continue
            qtd_projects += 1
            total_word_count += word_count
            dict_projects_text[file_name] = text
            st.write(f"📝 Processando arquivo: {file_name} ({word_count} palavras)")
            metadata["files"].append({"file_name": file_name, "word_count": word_count})

        # save metadata
        if qtd_projects > 0:
            with open(metadata_path, "w", encoding="utf-8") as f:
                json.dump(metadata, f, indent=4)

            # save files
            save_context_database(project_name, dict_projects_text, total_word_count)
            return True
        else:
            return False

    except Exception as e:
        st.error("Erro processando arquivos")
        return False


# Streamlit UI
st.title("Aplicações Cadastradas")

tab1, tab2, tab3, tab4 = st.tabs(["Exibir", "Inserir", "Atualizar", "Deletar"])

# Show Projects
with tab1:
    st.subheader("Projetos cadastrados")
    projects = [
        d
        for d in os.listdir(PROJECTS_DIR)
        if os.path.isdir(os.path.join(PROJECTS_DIR, d))
    ]

    if projects:
        list_projects = [""] + projects
        project_selected = st.selectbox("Selecione um projeto:", list_projects)
        if project_selected == "":
            st.write("")
            st.write("Por favor, selecione um projeto.")

        else:
            project_metadata_path = os.path.join(
                PROJECTS_DIR, project_selected, "metadata.json"
            )

            if os.path.exists(project_metadata_path):
                with open(project_metadata_path, "r", encoding="utf-8") as f:
                    metadata = json.load(f)

                st.write(f"**Nome do projeto:** {metadata['project_name']}")
                st.write(f"**Dono do Projeto:** {metadata['project_owner']}")
                st.write(f"**Data de criação:** {metadata['creation_date']}")
                st.write(f"**Arquivos submetidos:** {len(metadata['files'])}")

                if len(metadata["files"]) <= 10:

                    for file in metadata["files"]:
                        st.write(
                            f"- {file['file_name']} ({file['word_count']} palavras)"
                        )
            else:
                st.error("Não existem metadados para esse projeto.")

    else:
        st.write("Nenhum projeto foi cadastrado ainda.")

# Insert Project
with tab2:
    try:
        st.subheader("Criar um novo projeto")
        project_name = st.text_input("Nome do Projeto:", key="name_input_1")
        project_owner = st.text_input("Usuário do projeto:", key="user_input_1")
        password = st.text_input(
            "Senha do projeto (opcional):", type="password", key="password_input_1"
        )

        # application prompt
        with st.expander("Prompt (opcional)"):
            prompt_app = st.text_area(
                "Prompt do projeto:", height=200, key="prompt_input_1"
            )

        # Choose upload mode
        upload_mode = st.radio(
            "Como deseja inserir os documentos?",
            ["Upload manual", "SharePoint"],
            horizontal=True,
            key="radio_1",
        )

        # Manual Upload
        if upload_mode == "Upload manual":
            uploaded_files = st.file_uploader(
                "Upload de um Arquivo ou Zip:",
                type=[
                    "zip",
                    "rar",
                    "txt",
                    "csv",
                    "xlsx",
                    "docx",
                    "json",
                    "pdf",
                ],
                accept_multiple_files=True,
                key="uploader_1",
            )

        # Campos condicionais
        elif upload_mode == "SharePoint":
            sharepoint_url = st.text_input(
                "URL do SharePoint"
            )
            document_library = st.text_input(
                "Nome da biblioteca SharePoint"
            )
            sharepoint_url = sharepoint_url.strip()
            document_library = document_library.strip()

        else:
            uploaded_files = None

        # cast names
        try:
            project_name = project_name.strip()
            project_owner = project_owner.strip()
        except Exception:
            pass

        # Validate inputs
        flag_success = False
        if st.button("Upload", key="upload_1"):
            flag_valid, error_message = validate_project_inputs(
                project_name, project_owner, password, prompt_app, PROJECTS_DIR
            )
            if not flag_valid:
                st.error(error_message)
            else:
                # create project folder
                project_folder = os.path.join(PROJECTS_DIR, project_name)
                prompt_path = os.path.join(PROJECTS_DIR, project_name, "prompts")
                files_path = os.path.join(project_folder, "files")
                os.makedirs(files_path, exist_ok=True)

                if upload_mode == "Upload manual":

                    if uploaded_files is None or len(uploaded_files) == 0:
                        st.error("Por favor, envie um arquivo.")
                    else:

                        # Only one file
                        if len(uploaded_files) == 1:
                            file = uploaded_files[0]
                            if file.type in [
                                "application/zip",
                                "application/x-zip-compressed",
                            ]:

                                archive_path = os.path.join(project_folder, file.name)
                                with open(archive_path, "wb") as f:
                                    f.write(file.getbuffer())

                                with zipfile.ZipFile(archive_path, "r") as zip_ref:
                                    zip_ref.extractall(files_path)

                                os.remove(archive_path)

                            else:
                                file_path = os.path.join(files_path, file.name)

                                # Save file
                                with open(file_path, "wb") as f:
                                    f.write(file.getbuffer())

                        # Multiple files
                        else:
                            for file in uploaded_files:
                                file_path = os.path.join(files_path, file.name)

                                # Save file
                                with open(file_path, "wb") as f:
                                    f.write(file.getbuffer())

                else:
                    if sharepoint_url == "":
                        st.error("Por favor, preencha a URL corretamente.")
                        st.stop()
                    if document_library == "":
                        # default
                        document_library = "Português (Brasil)"

                    parameters = st.session_state["parameters"]
                    TENANT_ID = parameters.get("TENANT_ID")
                    CLIENT_ID = parameters.get("CLIENT_ID")
                    CLIENT_SECRET = parameters.get("CLIENT_SECRET")
                    URL_PATH = parameters.get("URL_PATH")
                    SITE_URL = f"{URL_PATH}/{sharepoint_url}:"

                    try:
                        sp = SharePointDownloader(
                            TENANT_ID,
                            CLIENT_ID,
                            CLIENT_SECRET,
                            SITE_URL,
                            document_library,
                            files_path,
                            False,
                        )
                        with st.spinner(
                            "Processando arquivos, por favor aguarde o término."
                        ):
                            sp.download_files_recursive()
                            time.sleep(1)
                    except Exception as e:
                        st.error(f"Erro ao baixar arquivos do SharePoint: {e}")
                        shutil.rmtree(project_folder, onerror=on_rm_error)
                        time.sleep(3)
                        st.stop()

                # save prompt
                if prompt_app != "":
                    os.makedirs(prompt_path, exist_ok=True)
                    prompt_file = os.path.join(prompt_path, "prompt_app.txt")
                    with open(prompt_file, "w", encoding="utf-8") as f:
                        f.write(prompt_app)

                # count files
                file_count = len(
                    [
                        f
                        for f in os.listdir(files_path)
                        if os.path.isfile(os.path.join(files_path, f))
                    ]
                )

                if file_count == 0:
                    st.error("Nenhum arquivo encontrado na pasta.")
                    shutil.rmtree(project_folder, onerror=on_rm_error)
                    st.stop()
                elif file_count == 1:
                    st.write("Foi mapeado apenas um arquivo.")
                    frase = "arquivo"
                else:
                    st.write(f"Foram mapeados um total de {file_count} arquivos.")
                    frase = "arquivos"

                # processing files
                with st.spinner(f"Processando {frase}..."):
                    flag_success = process_files_and_store_embeddings(
                        project_name, project_owner, password
                    )

                parameters = st.session_state["parameters"]
                if not FLAG_SAVE_FILE and parameters["delete_files"]:
                    try:
                        # Loop through all entries in the folder
                        for filename in os.listdir(files_path):
                            file_path = os.path.join(files_path, filename)
                            try:
                                os.remove(file_path)
                            except Exception:
                                pass
                        # remove folder
                        shutil.rmtree(files_path, onerror=on_rm_error)
                    except Exception:
                        pass

                if flag_success:
                    st.success(f"Projeto '{project_name}' criado com sucesso.")
                    time.sleep(3)
                else:
                    st.error(
                        f"Projeto '{project_name}' não foi criado devido a problemas nos arquivos."
                    )
                    if os.path.exists(project_folder) and os.path.isdir(project_folder):
                        shutil.rmtree(project_folder, onerror=on_rm_error)
                    time.sleep(3)

                # update page
                st.rerun()

    except Exception as e:
        st.error(f"Erro processando arquivos:\n{e}")
        project_path = os.path.join(PROJECTS_DIR, project_name)

        try:
            shutil.rmtree(project_path, onerror=on_rm_error)
        except Exception:
            pass

# Update Project
with tab3:

    try:
        st.subheader("Atualizar projeto")
        parameters = st.session_state["parameters"]

        # list all projects
        projects = [
            d
            for d in os.listdir(PROJECTS_DIR)
            if os.path.isdir(os.path.join(PROJECTS_DIR, d))
        ]

        if not projects:
            st.write("Nenhum projeto foi cadastrado ainda.")
            st.stop()

        else:
            flag_success = False
            list_projects = [""] + projects
            project_selected = st.selectbox("Projeto:", list_projects)
            if project_selected == "":
                st.write("")
                st.write("Por favor, selecione um projeto.")

            else:
                flag_success = False
                project_metadata_path = os.path.join(
                    PROJECTS_DIR, project_selected, "metadata.json"
                )

                if os.path.exists(project_metadata_path):
                    with open(project_metadata_path, "r", encoding="utf-8") as f:
                        metadata = json.load(f)

                    st.write(f"**Nome do projeto:** {metadata['project_name']}")
                    st.write(f"**Dono do Projeto:** {metadata['project_owner']}")
                    st.write(f"**Data de criação:** {metadata['creation_date']}")

                    if len(metadata["files"]) <= 10:
                        st.write(f"**Arquivos submetidos:** {len(metadata['files'])}")
                        for file in metadata["files"]:
                            st.write(
                                f"- {file['file_name']} ({file['word_count']} palavras)"
                            )

                    flag_success = True

                else:
                    st.error("Não existem metadados para esse projeto.")
                    st.stop()

            if flag_success:

                project_name = metadata["project_name"]
                project_owner = metadata["project_owner"]
                flag_password = metadata["flag_password"]
                project_password = metadata["password"].strip()
                project_password = decrypt_password(project_password).strip()

                if flag_password:
                    password = st.text_input(
                        "Senha do projeto:", type="password", key="password_input_2"
                    )
                else:
                    password = ""

                # Choose upload mode
                upload_mode = st.radio(
                    "Como deseja inserir os documentos?",
                    ["Upload manual", "SharePoint"],
                    horizontal=True,
                    key="radio_2",
                )

                # Manual Upload
                if upload_mode == "Upload manual":
                    uploaded_files = st.file_uploader(
                        "Upload de um Arquivo ou Zip:",
                        type=[
                            "zip",
                            "rar",
                            "txt",
                            "csv",
                            "xlsx",
                            "docx",
                            "json",
                            "pdf",
                        ],
                        accept_multiple_files=True,
                        key="uploader_2",
                    )

                # Campos condicionais
                elif upload_mode == "SharePoint":
                    sharepoint_url = st.text_input(
                        "URL do SharePoint"
                    )
                    document_library = st.text_input(
                        "Nome da biblioteca SharePoint [ex: Português (Brasil)]"
                    )
                    sharepoint_url = sharepoint_url.strip()
                    document_library = document_library.strip()

                else:
                    uploaded_files = None

                # cast names
                try:
                    project_name = project_name.strip()
                    project_owner = project_owner.strip()
                except Exception:
                    pass

                if st.button("Upload", key="upload_2"):
                    password = password.strip()
                    PASSWORD_ADMIN = parameters["admin"]
                    if (
                        flag_password
                        and password != project_password
                        and password != PASSWORD_ADMIN
                    ):
                        st.error("Senha incorreta.")
                        st.stop()

                    else:
                        # project folder
                        project_folder = os.path.join(PROJECTS_DIR, project_name)
                        files_path = os.path.join(project_folder, "files")

                        # Backup project folder if it exists
                        backup_dir = None
                        if os.path.exists(project_folder):
                            try:
                                backup_dir = tempfile.mkdtemp(prefix="project_backup_")
                                backup_path = os.path.join(
                                    backup_dir, os.path.basename(project_folder)
                                )
                                shutil.copytree(project_folder, backup_path)
                                shutil.rmtree(project_folder, onerror=on_rm_error)
                            except Exception as e:
                                st.error(f"Erro ao criar backup: {e}")
                                st.stop()

                        # create folders
                        os.makedirs(project_folder, exist_ok=True)
                        os.makedirs(files_path, exist_ok=True)

                        if upload_mode == "Upload manual":

                            if uploaded_files is None or len(uploaded_files) == 0:
                                st.error("Por favor, envie um arquivo.")
                            else:

                                # Only one file
                                if len(uploaded_files) == 1:
                                    file = uploaded_files[0]
                                    if file.type in [
                                        "application/zip",
                                        "application/x-zip-compressed",
                                    ]:

                                        archive_path = os.path.join(
                                            project_folder, file.name
                                        )
                                        with open(archive_path, "wb") as f:
                                            f.write(file.getbuffer())

                                        with zipfile.ZipFile(
                                            archive_path, "r"
                                        ) as zip_ref:
                                            zip_ref.extractall(files_path)

                                        os.remove(archive_path)

                                    else:
                                        file_path = os.path.join(files_path, file.name)

                                        # Save file
                                        with open(file_path, "wb") as f:
                                            f.write(file.getbuffer())

                                # Multiple files
                                else:
                                    for file in uploaded_files:
                                        file_path = os.path.join(files_path, file.name)

                                        # Save file
                                        with open(file_path, "wb") as f:
                                            f.write(file.getbuffer())

                        else:
                            if sharepoint_url == "":
                                st.error("Por favor, preencha a URL corretamente.")
                                st.stop()
                            if document_library == "":
                                # default
                                document_library = "Português (Brasil)"

                            parameters = st.session_state["parameters"]
                            TENANT_ID = parameters.get("TENANT_ID")
                            CLIENT_ID = parameters.get("CLIENT_ID")
                            CLIENT_SECRET = parameters.get("CLIENT_SECRET")
                            URL_PATH = parameters.get("URL_PATH")
                            SITE_URL = f"{URL_PATH}/{sharepoint_url}:"

                            try:
                                sp = SharePointDownloader(
                                    TENANT_ID,
                                    CLIENT_ID,
                                    CLIENT_SECRET,
                                    SITE_URL,
                                    document_library,
                                    files_path,
                                    False,
                                )
                                with st.spinner(
                                    "Processando arquivos, por favor aguarde o término."
                                ):
                                    sp.download_files_recursive()
                                    time.sleep(1)
                            except Exception as e:
                                st.error(f"Erro ao baixar arquivos do SharePoint: {e}")
                                shutil.rmtree(project_folder, onerror=on_rm_error)
                                time.sleep(3)
                                st.stop()

                        # count files
                        file_count = len(
                            [
                                f
                                for f in os.listdir(files_path)
                                if os.path.isfile(os.path.join(files_path, f))
                            ]
                        )

                        if file_count == 0:
                            st.error("Nenhum arquivo encontrado na pasta.")
                            shutil.rmtree(project_folder, onerror=on_rm_error)
                            st.stop()
                        elif file_count == 1:
                            st.write("Foi mapeado apenas um arquivo.")
                            frase = "arquivo"
                        else:
                            st.write(
                                f"Foram mapeados um total de {file_count} arquivos."
                            )
                            frase = "arquivos"

                        # processing files
                        with st.spinner(f"Processando {frase}..."):
                            flag_success = process_files_and_store_embeddings(
                                project_name, project_owner, project_password
                            )

                        # Save file
                        parameters = st.session_state["parameters"]
                        if not FLAG_SAVE_FILE and parameters["delete_files"]:
                            try:
                                # Loop through all entries in the folder
                                for filename in os.listdir(files_path):
                                    file_path = os.path.join(files_path, filename)
                                    try:
                                        os.remove(file_path)
                                    except Exception:
                                        pass
                                # remove folder
                                shutil.rmtree(files_path, onerror=on_rm_error)
                            except Exception:
                                pass

                        if flag_success:
                            st.success(
                                f"Projeto '{project_name}' atualizado com sucesso."
                            )
                            time.sleep(3)
                        else:
                            st.error(
                                f"Projeto '{project_name}' não foi criado devido a problemas nos arquivos."
                            )
                            if os.path.exists(project_folder) and os.path.isdir(
                                project_folder
                            ):
                                shutil.rmtree(project_folder, onerror=on_rm_error)
                            time.sleep(3)

                        try:
                            shutil.rmtree(backup_dir, onerror=on_rm_error)
                        except Exception:
                            pass

                        # update page
                        st.rerun()

    except Exception as e:
        st.error(f"Erro processando arquivos: {e}")
        project_path = os.path.join(PROJECTS_DIR, project_name)

        # Rollback if backup exists
        if backup_dir and os.path.exists(backup_path):
            try:
                if os.path.exists(project_path):
                    shutil.rmtree(project_path, onerror=on_rm_error)
                shutil.copytree(backup_path, project_path)
                st.warning("O projeto foi restaurado do backup temporário.")
            except Exception:
                pass
        else:
            try:
                shutil.rmtree(project_path, onerror=on_rm_error)
            except Exception:
                pass

        try:
            shutil.rmtree(backup_dir, onerror=on_rm_error)
        except Exception:
            pass

# Delete Project
with tab4:
    st.subheader("Deletar projeto")

    # Project list
    projects = [
        d
        for d in os.listdir(PROJECTS_DIR)
        if os.path.isdir(os.path.join(PROJECTS_DIR, d))
    ]

    if projects:
        list_projects = [""] + projects
        project_to_delete = st.selectbox(
            "Selecione um projeto para deletar:", list_projects
        )

        if project_to_delete == "":
            st.write("")
            st.write("Por favor, selecione um projeto.")

        else:

            # paths
            project_path = os.path.join(PROJECTS_DIR, project_to_delete)
            metadata_path = os.path.join(project_path, "metadata.json")

            # PASS
            parameters = st.session_state["parameters"]
            PASSWORD_ADMIN = parameters["admin"]

            # read project password
            with open(metadata_path, "r", encoding="utf-8") as f:
                metadata = json.load(f)
                flag_password = metadata["flag_password"]
                project_password = metadata["password"].strip()
                project_password = decrypt_password(project_password).strip()
                password = password.strip()

            # project password
            project_password = metadata["password"].strip()
            project_password = decrypt_password(project_password).strip()

            # password
            if flag_password:
                password = st.text_input(
                    "Senha do projeto:", type="password", key="password_input_3"
                )
            else:
                password = ""

            # delete button
            if st.button("Deletar"):
                if flag_password and password == "":
                    st.error("Digite uma senha.")
                else:
                    # check password
                    if password == project_password or password == PASSWORD_ADMIN:
                        # delete project
                        shutil.rmtree(
                            os.path.join(PROJECTS_DIR, project_to_delete),
                            onerror=on_rm_error,
                        )
                        st.success(
                            f"Projeto '{project_to_delete}' deletado com sucesso."
                        )
                        time.sleep(2)
                        st.rerun()
                    else:
                        st.error("Senha incorreta.")

    else:
        st.write("Nenhum projeto foi cadastrado ainda.")
