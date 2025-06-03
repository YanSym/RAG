from cryptography.fernet import Fernet
from langchain.schema import Document
from openai import OpenAI
import odf.opendocument
import streamlit as st
from io import BytesIO
import pandas as pd
import numpy as np
import subprocess
import platform
import odf.text
import zipfile
import PyPDF2
import base64
import time
import docx
import yaml
import json
import cv2
import re
import os

# projects dir
if platform.system() == "Windows":
    PROJECTS_DIR = "projects"
else:
    PROJECTS_DIR = os.path.join(os.getcwd(), "projects")

# Set environment variable explicitly for Tesseract (make sure it's correct for Windows)
if platform.system() == "Windows":
    TESSERACT_CMD = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
else:
    TESSERACT_CMD = "/usr/bin/tesseract"
    os.environ["TESSDATA_PREFIX"] = "/usr/share/tesseract-ocr/4.00/"


def initialize_application():
    if "flag_init" not in st.session_state or not st.session_state["flag_init"]:
        os.environ["STREAMLIT_WATCH_SYSTEM_FILES"] = "false"

        # read application parameters
        data = read_yaml_file("parameters/artifacts.yaml")

        # Flatten the dictionary by taking only the last level keys and values
        parameters = {
            key: value for subdict in data.values() for key, value in subdict.items()
        }

        # Store in Streamlit session state
        st.session_state["parameters"] = parameters
        st.session_state["chatbots"] = {}

        # set configs
        st.set_page_config(
            page_title="Aplicação RAG", page_icon="images/icon_rag.png"
        )

        # init
        st.session_state["flag_init"] = True


def read_yaml_file(file_path):
    """Reads a YAML file and returns its contents as a dictionary."""
    with open(file_path, "r", encoding="utf-8") as file:
        data = yaml.safe_load(file)
    return data


def get_llm_response(prompt: str, parameters: dict, temperature: float = 0) -> str:
    """
    Sends a prompt to the OpenAI API and returns the response.

    Args:
        prompt (str): The prompt to send to the API.
        parameters (dict): The parameters to use for the API request.
        temperature (float): The temperature setting for the API request.

    Returns:
        str: The response from the OpenAI API.
    """
    client = OpenAI(api_key=parameters["key"])
    response = client.chat.completions.create(
        model=parameters["model_name"],
        temperature=temperature,
        messages=[{"role": "user", "content": prompt}],
    )
    response = response.choices[0].message.content
    return response


def show_sidebar():
    # Add logo and title at the top of the sidebar
    st.sidebar.markdown("<br>" * 12, unsafe_allow_html=True)

    # Custom CSS to push the image to the bottom
    st.markdown(
        """
        <style>
            [data-testid="stSidebar"] > div:first-child {
                display: flex;
                flex-direction: column;
                height: 100vh;
            }
            [data-testid="stSidebar"] > div:first-child > div:nth-child(2) {
                margin-top: auto;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Display the image at the bottom
    st.sidebar.image("images/logo_rag.png", width=150)


def load_documents(directory):
    # iterate directory
    documents = []
    for file in os.listdir(directory):
        try:
            file_path = os.path.join(directory, file)
            if os.path.isfile(file_path) and file.split(".")[-1].lower() in [
                "txt",
                "xlsx",
                "csv",
                "json",
                "yaml",
                "yml",
                "odt",
                "docx",
                "pdf",
            ]:
                text = extract_text_from_file(file_path)
                if text:
                    documents.append(
                        Document(content=text, metadata={"file_name": file})
                    )

        except Exception as e:
            pass

    if len(documents) == 0:
        documents = [""]

    return documents


def encrypt_password(password: str) -> str:
    """Encrypt and store the password in a Base64-encoded format."""
    try:
        parameters = st.session_state["parameters"]
        key = parameters["secret"]
        fernet = Fernet(key)

        # Encrypt the password
        encrypted_password = fernet.encrypt(password.encode())

        # Encode in Base64 and return
        encrypted_password_b64 = base64.b64encode(encrypted_password).decode("utf-8")
        return encrypted_password_b64
    except Exception as e:
        st.error(f"Error encrypting password: {str(e)}")
        return ""


def decrypt_password(password_b64: str) -> str:
    """Retrieve and decrypt the Base64-encoded password."""
    try:
        parameters = st.session_state["parameters"]
        key = parameters["secret"]
        fernet = Fernet(key)

        # Ensure proper Base64 padding (if needed)
        missing_padding = len(password_b64) % 4
        if missing_padding:
            password_b64 += "=" * (4 - missing_padding)

        # Decode the Base64 string
        encrypted_password = base64.b64decode(password_b64.encode("utf-8"))
        decrypted_password = fernet.decrypt(encrypted_password).decode("utf-8")
        return decrypted_password

    except Exception as e:
        return ""


def check_os():
    system_name = platform.system()
    if system_name == "Windows":
        return "Windows"
    elif system_name == "Linux":
        return "Linux"
    else:
        return "Other"


def preprocess_image(image_path: str) -> np.ndarray:
    """Preprocess the image for better OCR results."""
    image = cv2.imread(image_path)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)  # Convert to grayscale
    gray = cv2.GaussianBlur(gray, (5, 5), 0)  # Apply Gaussian blur
    _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    return thresh


def extract_text_from_image(image_path):
    try:
        # Preprocess the image before running OCR
        processed_img = preprocess_image(image_path)

        # Save the preprocessed image temporarily to run tesseract on it
        temp_image_path = (
            "/tmp/preprocessed_image.jpg"
            if platform.system() != "Windows"
            else "C:/temp/preprocessed_image.jpg"
        )
        cv2.imwrite(temp_image_path, processed_img)

        # Define the subprocess command to run tesseract OCR
        result = subprocess.run(
            [TESSERACT_CMD, temp_image_path, "stdout"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )

        if result.stderr:
            st.write(f"Erro na leitura da imagem: {result.stderr}")
            os.remove(temp_image_path)
            time.sleep(2)
            return None

        # Return result
        else:
            return result.stdout.strip()
    except Exception as e:
        st.write("Erro processando imagem:", str(e))
        time.sleep(2)
        return None


def process_uploaded_file(uploaded_file, temp_dir):
    """
    Extracts the contents of a zip file or processes an individual document.

    Args:
        uploaded_file: The uploaded file (zip or document).
        temp_dir: Temporary directory to store extracted files (for zip).

    Returns:
        List[str]: List of file paths to process.
    """
    files = []

    # If it's a zip file, extract all contents
    if uploaded_file.type in ["application/zip", "application/x-zip-compressed"]:
        with zipfile.ZipFile(uploaded_file, "r") as zip_ref:
            zip_ref.extractall(temp_dir)
            files = [os.path.join(temp_dir, f) for f in os.listdir(temp_dir)]
    else:
        # If it's an individual file, just save it temporarily
        temp_file_path = os.path.join(temp_dir, uploaded_file.name)
        with open(temp_file_path, "wb") as f:
            f.write(uploaded_file.read())
        files.append(temp_file_path)

    return files


# Function to extract text from different file types
def extract_text_from_file(file_path):
    try:
        ext = file_path.split(".")[-1].lower()
        text = ""

        # text
        if ext in ["txt"]:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

        # csv file
        elif ext in ["csv", "xlsx"]:
            df = pd.read_csv(file_path) if ext == "csv" else pd.read_excel(file_path)
            text = df.to_string()

        elif ext in ["json"]:
            with open(file_path, "r", encoding="utf-8") as f:
                text = json.dumps(json.load(f))

        elif ext in ["yaml", "yml"]:
            with open(file_path, "r", encoding="utf-8") as f:
                text = yaml.dump(yaml.safe_load(f))

        elif ext in ["docx"]:
            doc = docx.Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])

        elif ext in ["odt"]:
            odt_doc = odf.opendocument.load(file_path)
            text = "\n".join([el.text for el in odt_doc.getElementsByType(odf.text.P)])

        elif ext in ["pdf"]:
            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = "\n".join(
                    [
                        pdf_reader.pages[i].extract_text()
                        for i in range(len(pdf_reader.pages))
                    ]
                )

        if text:
            text = re.sub(r"\s+", " ", text)
            if len(text) <= 5:
                return None
            return text.strip()

        return None

    except Exception:
        return None


def check_run_docker():
    try:
        with open("/proc/1/cgroup", "rt") as f:
            return any("docker" in line for line in f)
    except FileNotFoundError:
        return False


def convert_data(data: list) -> pd.DataFrame:
    """
    Converts the extracted data into a Pandas DataFrame
    and sorts it by candidate evaluation and estimated salary.

    Args:
        data (list): A list of dictionaries containing candidate information.

    Returns:
        pd.DataFrame: A sorted DataFrame with cleaned data.
    """
    df = pd.DataFrame(data)
    df["Nome"] = df["Nome"].astype(str)
    df["Nome"] = df["Nome"].str.title().str.rstrip()
    df["Telefone"] = df["Telefone"].astype(str)
    try:
        df["Telefone"] = (
            df["Telefone"].apply(lambda x: int(re.sub(r"\D", "", x))).astype(str)
        )
    except Exception:
        pass

    df["E-mail"] = df["E-mail"].apply(lambda x: x if ".com" in x.lower() else "")
    df["Linkedin"] = df["Linkedin"].apply(
        lambda x: x if "linkedin" in x.lower() else ""
    )
    df["Linkedin"] = df["Linkedin"].str.replace("www.", "")
    df["Linkedin"] = df["Linkedin"].str.rstrip("/")
    df["Git"] = df["Git"].str.replace("https://", "")
    df["Git"] = df["Git"].str.replace("www.", "")
    df["Git"] = df["Git"].str.rstrip("/")
    df["Nome"] = df["Nome"].str.title().str.strip()
    df["Localização"] = df["Localização"].str.title().str.strip()
    df["Cargo Atual"] = df["Cargo Atual"].str.title().str.strip()
    df["Escolaridade"] = df["Escolaridade"].str.title().str.strip()
    df["Escola"] = df["Escola"].str.title().str.strip()
    df["Empresa"] = df["Empresa"].str.title().str.strip()
    df["Anos de Experiência"] = df["Anos de Experiência"].astype(int)
    df["Habilidades"] = df["Habilidades"].astype(str)
    df["Habilidades"] = df["Habilidades"].str.strip()
    df["Empresa"] = df["Empresa"].str.rstrip()
    df["Avaliação do Candidato"] = pd.to_numeric(
        df["Avaliação do Candidato"], errors="coerce"
    )
    df["Salario Estimado"] = pd.to_numeric(df["Salario Estimado"], errors="coerce")
    df = df.sort_values(
        by=["Avaliação do Candidato", "Salario Estimado"],
        ascending=False,
    )
    df = df.loc[~df["Nome"].isnull()]
    df = df.reset_index(drop=True)
    return df


def download_excel(df: pd.DataFrame) -> BytesIO:
    """
    Converts a Pandas DataFrame to an Excel file and returns it as a BytesIO object for download.

    Args:
        df (pd.DataFrame): The DataFrame to convert.

    Returns:
        BytesIO: The Excel file as a BytesIO object.
    """
    excel_file = BytesIO()
    df.to_excel(excel_file, index=False)
    excel_file.seek(0)
    return excel_file
